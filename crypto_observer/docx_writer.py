from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .config import REPORT_TITLE, SECTION_ORDER
from .text_utils import normalize_chinese_punctuation

FONT_BODY = "仿宋"
FONT_TOC_ITEM = "楷体"
FONT_HEADING = "黑体"
FONT_LATIN = "Times New Roman"
HEADING_COLOR = "2F5496"

STYLE_ARTICLE_TITLE = "文章一级标题"
STYLE_BODY = "正文内容"
STYLE_SOURCE = "信息来源"
STYLE_REFERENCE = "原文信息"


def _ensure_style(doc: Document, name: str):
    styles = doc.styles
    if name in styles:
        return styles[name]
    return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _set_rfonts(obj, east_asia: str, ascii_font: str = FONT_LATIN) -> None:
    rpr = obj._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:cs"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), east_asia)


def _set_run_font(run, east_asia: str = FONT_BODY, size: float = 14, bold: bool = False, color: str | None = None) -> None:
    run.font.name = FONT_LATIN
    _set_rfonts(run, east_asia)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_style_font(style, east_asia: str, size: float, bold: bool | None = None, color: str | None = None) -> None:
    style.font.name = FONT_LATIN
    _set_rfonts(style, east_asia)
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def _set_doc_grid(section, line_pitch: int = 312) -> None:
    sect_pr = section._sectPr
    for child in sect_pr.findall(qn("w:docGrid")):
        sect_pr.remove(child)
    doc_grid = OxmlElement("w:docGrid")
    doc_grid.set(qn("w:type"), "lines")
    doc_grid.set(qn("w:linePitch"), str(line_pitch))
    sect_pr.append(doc_grid)


def _setup_section(section) -> None:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    section.header_distance = Cm(1.50)
    section.footer_distance = Cm(1.75)
    _set_doc_grid(section)


def _set_single_grid_paragraph(fmt, *, first_line: bool = False, justify: bool = False) -> None:
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.first_line_indent = Pt(28) if first_line else Pt(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = 1


def _setup_document(doc: Document) -> None:
    for section in doc.sections:
        _setup_section(section)

    normal = doc.styles["Normal"]
    _set_style_font(normal, FONT_BODY, 14)
    _set_single_grid_paragraph(normal.paragraph_format, first_line=False)

    heading1 = doc.styles["Heading 1"]
    _set_style_font(heading1, FONT_HEADING, 15, color=HEADING_COLOR)
    _set_single_grid_paragraph(heading1.paragraph_format, first_line=False)
    heading1.paragraph_format.space_before = Pt(3)
    heading1.paragraph_format.space_after = Pt(12)

    article_title = _ensure_style(doc, STYLE_ARTICLE_TITLE)
    _set_style_font(article_title, FONT_HEADING, 15)
    _set_single_grid_paragraph(article_title.paragraph_format, first_line=False)
    article_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    body = _ensure_style(doc, STYLE_BODY)
    body.base_style = normal
    _set_style_font(body, FONT_BODY, 14)
    _set_single_grid_paragraph(body.paragraph_format, first_line=True)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    source = _ensure_style(doc, STYLE_SOURCE)
    source.base_style = normal
    _set_style_font(source, FONT_BODY, 10.5)
    _set_single_grid_paragraph(source.paragraph_format, first_line=False)
    source.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    source.paragraph_format.space_before = Pt(2.5)
    source.paragraph_format.space_after = Pt(5)

    reference = _ensure_style(doc, STYLE_REFERENCE)
    reference.base_style = normal
    _set_style_font(reference, FONT_BODY, 9)
    _set_single_grid_paragraph(reference.paragraph_format, first_line=False)
    reference.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _apply_body_format(paragraph, *, first_line: bool = True, align: int = WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    _set_single_grid_paragraph(fmt, first_line=first_line)


def _add_paragraph(doc: Document, text: str, *, style: str | None = None, font: str = FONT_BODY, size: float = 14, bold: bool = False, align: int | None = None, color: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(normalize_chinese_punctuation(text))
    _set_run_font(r, east_asia=font, size=size, bold=bold, color=color)


def _items_by_section(items: list[dict[str, Any]], section: str) -> list[dict[str, Any]]:
    return [row for row in items if row.get("section") == section]


def _body_paragraphs(row: dict[str, Any]) -> list[str]:
    paragraphs = row.get("body_paragraphs")
    if isinstance(paragraphs, list):
        cleaned = [normalize_chinese_punctuation(str(x).strip()) for x in paragraphs if str(x).strip()]
    else:
        cleaned = []
    if not cleaned:
        for key in ("lead_cn", "summary_cn", "analysis_cn"):
            value = normalize_chinese_punctuation(str(row.get(key) or "").strip())
            if value and value != "-":
                cleaned.append(value)
    return cleaned or ["-"]


def _key_points(row: dict[str, Any]) -> list[str]:
    raw = row.get("key_points")
    if isinstance(raw, list):
        return [normalize_chinese_punctuation(str(x).strip()) for x in raw if str(x).strip()][:3]
    return []


def _write_toc(doc: Document, items: list[dict[str, Any]]) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目  录")
    _set_run_font(r, east_asia=FONT_HEADING, size=16)

    for section in SECTION_ORDER:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"【{section}】")
        _set_run_font(r, east_asia=FONT_HEADING, size=16, color=HEADING_COLOR)

        rows = _items_by_section(items, section)
        if not rows:
            _write_toc_item(doc, "暂无满足自动筛选条件的资讯")
            continue
        for row in rows:
            _write_toc_item(doc, str(row.get("title_cn") or "-"))


def _write_toc_item(doc: Document, title: str) -> None:
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(f"· {normalize_chinese_punctuation(title)}")
    _set_run_font(r, east_asia=FONT_TOC_ITEM, size=14)


def _write_section_heading(doc: Document, section: str) -> None:
    _add_paragraph(doc, f"【{section}】", style="Heading 1", font=FONT_HEADING, size=15, color=HEADING_COLOR)


def _write_article_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph(style=STYLE_ARTICLE_TITLE)
    _apply_body_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(normalize_chinese_punctuation(title))
    _set_run_font(r, east_asia=FONT_HEADING, size=15)


def _write_key_point(doc: Document, point: str) -> None:
    p = doc.add_paragraph()
    _apply_body_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    r = p.add_run(f"· {normalize_chinese_punctuation(point)}")
    _set_run_font(r, east_asia=FONT_BODY, size=14)


def _write_body(doc: Document, paragraph: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph(style=STYLE_BODY)
    _apply_body_format(p, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    r = p.add_run(normalize_chinese_punctuation(paragraph))
    _set_run_font(r, east_asia=FONT_BODY, size=14, bold=bold)


def _write_source(doc: Document, source_name: str) -> None:
    p = doc.add_paragraph(style=STYLE_SOURCE)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_single_grid_paragraph(p.paragraph_format, first_line=False)
    p.paragraph_format.space_before = Pt(2.5)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"（信息来源：{normalize_chinese_punctuation(source_name)}）")
    _set_run_font(r, east_asia=FONT_BODY, size=10.5)


def _write_reference(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style=STYLE_REFERENCE)
    _apply_body_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run(normalize_chinese_punctuation(text))
    _set_run_font(r, east_asia=FONT_BODY, size=9)


def _write_article(doc: Document, row: dict[str, Any]) -> None:
    title = str(row.get("title_cn") or "-").strip()
    _write_article_title(doc, title)

    for point in _key_points(row):
        _write_key_point(doc, point)

    lead = normalize_chinese_punctuation(str(row.get("lead_cn") or "").strip())
    if lead and lead != "-":
        _write_body(doc, lead)

    for paragraph in _body_paragraphs(row):
        if paragraph == lead:
            continue
        _write_body(doc, paragraph)

    source_name = str(row.get("source_name") or "-").strip()
    source_title = str(row.get("source_title") or "-").strip()
    url = str(row.get("url") or "-").strip()
    _write_source(doc, source_name)
    if source_title and source_title != "-":
        _write_reference(doc, f"原文标题：{source_title}")
    if url and url != "-":
        _write_reference(doc, f"原文链接：{url}")


def write_docx(report: dict[str, Any], output_path: str | Path, metadata: dict[str, Any] | None = None) -> Path:
    metadata = metadata or {}
    items = [x for x in report.get("items", []) if isinstance(x, dict)]
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _setup_document(doc)
    _write_toc(doc, items)

    first_section = True
    for section in SECTION_ORDER:
        if first_section:
            body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
            _setup_section(body_section)
        else:
            next_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
            _setup_section(next_section)
        first_section = False
        _write_section_heading(doc, section)
        rows = _items_by_section(items, section)
        if not rows:
            _write_body(doc, "暂无满足自动筛选条件的资讯。")
            continue
        for idx, row in enumerate(rows, start=1):
            _write_article(doc, row)
            if idx != len(rows):
                doc.add_paragraph()

    doc.core_properties.title = f"{REPORT_TITLE}周刊"
    doc.core_properties.subject = str(metadata.get("period") or "最近三天")
    doc.save(path)
    return path
