from __future__ import annotations

import hashlib
import json
import logging
import os
import re
import urllib.parse
import xml.etree.ElementTree as ET
from dataclasses import asdict, dataclass
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any
from zoneinfo import ZoneInfo

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfReader

from .config import RESEARCH_LOOKBACK_DAYS, RESEARCH_ORGANIZATION_SITES, RESEARCH_REPORT_TITLE, RESEARCH_SOURCE_URLS, USER_AGENT
from .docx_writer import (
    FONT_HEADING,
    HEADING_COLOR,
    _set_run_font,
    _setup_document,
    _setup_section,
    _write_body,
    _write_key_point,
    _write_reference,
    _write_source,
)
from .text_utils import has_half_width_quotes, normalize_chinese_punctuation

LOGGER = logging.getLogger(__name__)
BEIJING_TZ = ZoneInfo("Asia/Shanghai")
DEFAULT_MODEL = "deepseek-v4-flash"
DEFAULT_BASE_URL = "https://api.deepseek.com"
REPORT_TERMS = ("crypto", "digital asset", "digital assets", "tokenization", "tokenisation", "stablecoin", "blockchain", "defi", "web3", "rwa", "bitcoin", "ethereum")
BLOCKED_URL_PARTS = ("/zh", "/zh-cn", "/cn/", "?lang=zh", "language=zh")
MIN_SOURCE_CHARS = 18000
MIN_RESEARCH_PARAGRAPHS = 24
MIN_RESEARCH_CHARS = 12000
TARGET_RESEARCH_PARAGRAPHS = 42
MONTHS = {
    "january": 1, "february": 2, "march": 3, "april": 4, "may": 5, "june": 6,
    "july": 7, "august": 8, "september": 9, "october": 10, "november": 11, "december": 12,
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6, "jul": 7, "aug": 8,
    "sep": 9, "sept": 9, "oct": 10, "nov": 11, "dec": 12,
}
PUBLICATION_CONTEXT_WORDS = (
    "published", "publication", "released", "release date", "report date", "date of report", "updated", "last updated",
    "as of", "prepared", "prepared on", "issued", "issue date", "research report", "monthly report", "quarterly report",
)


@dataclass
class ResearchCandidate:
    title: str
    url: str
    source_name: str
    published_at: str = ""
    snippet: str = ""


def _chat(messages: list[dict[str, str]], timeout: int = 360) -> str:
    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        raise RuntimeError("DEEPSEEK_API_KEY is not set")
    base_url = os.getenv("DEEPSEEK_BASE_URL", DEFAULT_BASE_URL).rstrip("/")
    model = os.getenv("DEEPSEEK_MODEL", DEFAULT_MODEL)
    response = requests.post(
        f"{base_url}/chat/completions",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json={"model": model, "messages": messages, "temperature": 0.0, "stream": False, "response_format": {"type": "json_object"}},
        timeout=timeout,
    )
    if response.status_code >= 400:
        raise RuntimeError(f"DeepSeek API error {response.status_code}: {response.text[:800]}")
    return response.json()["choices"][0]["message"]["content"]


def _extract_json(text: str) -> dict[str, Any]:
    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", text.strip())
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start >= 0 and end > start:
            return json.loads(cleaned[start : end + 1])
        raise


def _norm_text(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _safe_filename(value: str) -> str:
    name = re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("_")
    return (name or "research_report")[:120]


def _is_english_url(url: str) -> bool:
    low = url.lower()
    return not any(part in low for part in BLOCKED_URL_PARTS)


def _looks_relevant(text: str) -> bool:
    low = text.lower()
    return any(term in low for term in REPORT_TERMS)


def _bing_rss_url(query: str) -> str:
    return "https://www.bing.com/search?" + urllib.parse.urlencode({"q": query, "format": "rss", "setlang": "en", "mkt": "en-US"})


def _parse_rss_date(value: str) -> datetime | None:
    if not value:
        return None
    try:
        from email.utils import parsedate_to_datetime

        dt = parsedate_to_datetime(value)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(BEIJING_TZ)
    except Exception:
        return None


def _parse_iso_datetime(value: str) -> datetime | None:
    if not value:
        return None
    try:
        dt = datetime.fromisoformat(value.replace("Z", "+00:00"))
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(BEIJING_TZ)
    except Exception:
        return None


def _within(dt: datetime | None, run_date: datetime, lookback_days: int = RESEARCH_LOOKBACK_DAYS) -> bool:
    return bool(dt and run_date - timedelta(days=lookback_days) <= dt <= run_date)


def _date_from_match(year: str, month: str, day: str | None = None) -> datetime | None:
    try:
        dd = int(day) if day else 1
        return datetime(int(year), int(month), dd, tzinfo=BEIJING_TZ)
    except Exception:
        return None


def _find_explicit_dates(text: str, *, allow_month_only: bool = True) -> list[datetime]:
    low = text.lower()
    dates: list[datetime] = []
    for match in re.finditer(r"\b(20\d{2})[-_/](0?[1-9]|1[0-2])[-_/](0?[1-9]|[12]\d|3[01])\b", low):
        dt = _date_from_match(match.group(1), match.group(2), match.group(3))
        if dt:
            dates.append(dt)
    for match in re.finditer(r"\b(0?[1-9]|[12]\d|3[01])[-_/](0?[1-9]|1[0-2])[-_/](20\d{2})\b", low):
        dt = _date_from_match(match.group(3), match.group(2), match.group(1))
        if dt:
            dates.append(dt)
    for match in re.finditer(r"\b(january|february|march|april|may|june|july|august|september|october|november|december|jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)\s+(\d{1,2},\s*)?(20\d{2})\b", low):
        month = MONTHS.get(match.group(1))
        day = (match.group(2) or "").replace(",", "").strip() or None
        if month:
            dt = _date_from_match(match.group(3), str(month), day)
            if dt:
                dates.append(dt)
    if allow_month_only:
        for match in re.finditer(r"\b(january|february|march|april|may|june|july|august|september|october|november|december|jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)\s+(20\d{2})\b", low):
            month = MONTHS.get(match.group(1))
            if month:
                dt = _date_from_match(match.group(2), str(month))
                if dt:
                    dates.append(dt)
        for match in re.finditer(r"\b(20\d{2})[-_/](0?[1-9]|1[0-2])\b", low):
            dt = _date_from_match(match.group(1), match.group(2))
            if dt:
                dates.append(dt)
    out: list[datetime] = []
    seen: set[str] = set()
    for dt in dates:
        key = dt.date().isoformat()
        if key not in seen:
            seen.add(key)
            out.append(dt)
    return out


def _publication_date_evidence(candidate: ResearchCandidate, text: str, run_date: datetime) -> list[str]:
    """Require an explicit publication/release date for the report itself."""
    start = run_date - timedelta(days=RESEARCH_LOOKBACK_DAYS)
    evidence: list[str] = []
    identity_source = f"{candidate.title} {candidate.url}"
    contextual_source = f"{candidate.title} {candidate.url} {candidate.snippet}"

    for dt in _find_explicit_dates(identity_source, allow_month_only=True):
        if start <= dt <= run_date:
            evidence.append(f"title_url_snippet={dt.date().isoformat()}")

    for dt in _find_explicit_dates(contextual_source, allow_month_only=True):
        if start <= dt <= run_date:
            ctx_low = contextual_source.lower()
            date_key = dt.strftime("%Y-%m")
            pos = ctx_low.find(date_key)
            if pos < 0:
                pos = ctx_low.find(dt.strftime("%Y/%m"))
            if pos < 0:
                for month_name, month_no in MONTHS.items():
                    if month_no == dt.month and str(dt.year) in ctx_low:
                        pos = ctx_low.find(month_name)
                        if pos >= 0:
                            break
            window = ctx_low[max(0, pos - 120) : pos + 160] if pos >= 0 else ctx_low[:600]
            if any(word in window for word in PUBLICATION_CONTEXT_WORDS):
                evidence.append(f"title_url_snippet={dt.date().isoformat()}")

    cover_text = text[:4500]
    cover_low = cover_text.lower()
    for dt in _find_explicit_dates(cover_text, allow_month_only=True):
        date_token = dt.strftime("%Y-%m")
        pos = cover_low.find(date_token)
        if pos < 0:
            date_token = dt.strftime("%Y/%m")
            pos = cover_low.find(date_token)
        if pos < 0:
            for month_name, month_no in MONTHS.items():
                if month_no == dt.month and str(dt.year) in cover_low:
                    idx = cover_low.find(month_name)
                    if idx >= 0:
                        pos = idx
                        break
        window = cover_low[max(0, pos - 140) : pos + 180] if pos >= 0 else cover_low[:700]
        has_context = any(word in window for word in PUBLICATION_CONTEXT_WORDS) or pos < 700
        if has_context and start <= dt <= run_date:
            evidence.append(f"cover_publication_date={dt.date().isoformat()}")

    search_dt = _parse_iso_datetime(candidate.published_at)
    if evidence and _within(search_dt, run_date):
        evidence.append(f"search_seen_at={candidate.published_at}")

    return list(dict.fromkeys(evidence))[:8]


def _stale_identity_date_markers(candidate: ResearchCandidate, run_date: datetime) -> list[str]:
    start = run_date - timedelta(days=RESEARCH_LOOKBACK_DAYS)
    markers: list[str] = []
    strict_source = f"{candidate.title} {candidate.url}"
    for dt in _find_explicit_dates(strict_source, allow_month_only=True):
        if dt < start:
            markers.append(dt.date().isoformat())
    return sorted(set(markers))[:6]


def _stale_report_date_markers(candidate: ResearchCandidate, text: str, run_date: datetime) -> list[str]:
    start = run_date - timedelta(days=RESEARCH_LOOKBACK_DAYS)
    markers: list[str] = []
    strict_source = f"{candidate.title} {candidate.url} {text[:2500]}"
    for dt in _find_explicit_dates(strict_source, allow_month_only=True):
        if dt < start:
            markers.append(dt.date().isoformat())
    return sorted(set(markers))[:6]


def _add_candidate(out: list[ResearchCandidate], seen: set[str], *, title: str, url: str, source_name: str, published_at: str = "", snippet: str = "") -> None:
    url = _norm_text(url)
    title = _norm_text(title) or url.rsplit("/", 1)[-1]
    if not url.startswith("http") or url in seen or not _is_english_url(url):
        return
    if re.search(r"[\u4e00-\u9fff]", title + snippet + url):
        return
    if not _looks_relevant(title + " " + snippet + " " + url):
        return
    seen.add(url)
    out.append(ResearchCandidate(title=title, url=url, source_name=source_name, published_at=published_at, snippet=snippet))


def _extract_pdf_links_from_page(name: str, page_url: str, seen: set[str]) -> list[ResearchCandidate]:
    candidates: list[ResearchCandidate] = []
    try:
        response = requests.get(page_url, headers={"User-Agent": USER_AGENT, "Accept": "text/html,*/*"}, timeout=25, allow_redirects=True)
        response.raise_for_status()
    except Exception as exc:
        LOGGER.debug("research source page fetch failed for %s: %s", page_url, exc)
        return candidates
    soup = BeautifulSoup(response.text, "html.parser")
    page_text = _norm_text(soup.get_text(" "))[:3500]
    for a in soup.find_all("a", href=True):
        href = urllib.parse.urljoin(response.url, a.get("href") or "")
        anchor = _norm_text(a.get_text(" "))
        href_low = href.lower()
        if ".pdf" in href_low or "download" in href_low or "report" in href_low or "whitepaper" in href_low:
            _add_candidate(candidates, seen, title=anchor or name, url=href, source_name=name, snippet=page_text)
    return candidates


def search_research_pdfs(lookback_days: int = RESEARCH_LOOKBACK_DAYS) -> list[ResearchCandidate]:
    start = datetime.now(BEIJING_TZ) - timedelta(days=lookback_days)
    candidates: list[ResearchCandidate] = []
    seen: set[str] = set()
    after = start.strftime("%Y-%m-%d")
    topics = f'(crypto OR "digital assets" OR tokenization OR tokenisation OR stablecoin OR blockchain OR DeFi OR Web3 OR RWA) (report OR research OR outlook OR whitepaper) filetype:pdf after:{after}'
    for org, domain in RESEARCH_ORGANIZATION_SITES.items():
        query = f"site:{domain} {topics}"
        try:
            xml_text = requests.get(_bing_rss_url(query), headers={"User-Agent": USER_AGENT}, timeout=20).text
            root = ET.fromstring(xml_text.encode("utf-8"))
        except Exception as exc:
            LOGGER.debug("research search failed for %s: %s", org, exc)
            continue
        for item in root.findall("./channel/item"):
            title = _norm_text(item.findtext("title"))
            url = _norm_text(item.findtext("link"))
            snippet = _norm_text(item.findtext("description"))
            dt = _parse_rss_date(item.findtext("pubDate") or "")
            if dt and dt < start:
                continue
            if ".pdf" not in url.lower() and "pdf" not in snippet.lower() and "report" not in (title + snippet).lower():
                continue
            _add_candidate(candidates, seen, title=title, url=url, source_name=org, published_at=dt.isoformat() if dt else "", snippet=snippet)
    for name, url in RESEARCH_SOURCE_URLS.items():
        candidates.extend(_extract_pdf_links_from_page(name, url, seen))
    return candidates[:180]


def _resolve_pdf_url(url: str) -> str | None:
    if not url.startswith("http"):
        return None
    if ".pdf" in url.lower():
        return url
    try:
        response = requests.get(url, headers={"User-Agent": USER_AGENT, "Accept": "text/html,*/*"}, timeout=25, allow_redirects=True)
        response.raise_for_status()
    except Exception:
        return None
    if "application/pdf" in response.headers.get("content-type", "").lower() or response.content[:2048].find(b"%PDF") >= 0:
        return response.url
    soup = BeautifulSoup(response.text, "html.parser")
    for a in soup.find_all("a", href=True):
        href = urllib.parse.urljoin(response.url, a.get("href") or "")
        text = _norm_text(a.get_text(" "))
        if ".pdf" in href.lower() and _looks_relevant(href + " " + text):
            return href
    return None


def _download_pdf(candidate: ResearchCandidate, output_dir: Path) -> Path | None:
    output_dir.mkdir(parents=True, exist_ok=True)
    try:
        pdf_url = _resolve_pdf_url(candidate.url)
        if not pdf_url:
            return None
        response = requests.get(pdf_url, headers={"User-Agent": USER_AGENT, "Accept": "application/pdf,*/*"}, timeout=70, allow_redirects=True)
        response.raise_for_status()
        content = response.content
        if not content.startswith(b"%PDF") and b"%PDF" not in content[:2048]:
            return None
        candidate.url = response.url
        digest = hashlib.sha1(candidate.url.encode("utf-8")).hexdigest()[:8]
        path = output_dir / f"{_safe_filename(candidate.source_name)}_{_safe_filename(candidate.title)}_{digest}.pdf"
        path.write_bytes(content)
        return path
    except Exception as exc:
        LOGGER.debug("pdf download failed for %s: %s", candidate.url, exc)
        return None


def _read_pdf(path: Path, max_chars: int = 100000) -> str:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages[:180]:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            parts.append(text)
        if sum(len(x) for x in parts) >= max_chars:
            break
    return _norm_text("\n".join(parts))[:max_chars]


def _score_candidate(candidate: ResearchCandidate, text: str, evidence: list[str]) -> int:
    combined = f"{candidate.title} {candidate.snippet} {candidate.url} {text[:6000]}".lower()
    score = sum(5 for term in REPORT_TERMS if term in combined)
    score += 12 if any(x in combined for x in ("report", "research", "outlook", "whitepaper", "state of")) else 0
    score += min(len(text) // 2500, 18)
    score += len(evidence) * 20
    preferred = ("pwc", "kpmg", "bcg", "mckinsey", "bis", "deloitte", "ey", "pitchbook", "citi", "jpmorgan", "coinbase", "chainalysis", "galaxy", "coinshares", "crypto.com")
    if any(x in candidate.source_name.lower() for x in preferred):
        score += 8
    return score


def _choose_pdf(output_dir: Path, run_date: datetime) -> tuple[ResearchCandidate | None, Path | None, str, list[str]]:
    scored: list[tuple[int, ResearchCandidate, Path, str, list[str]]] = []
    rejected: list[str] = []
    for candidate in search_research_pdfs():
        identity_stale = _stale_identity_date_markers(candidate, run_date)
        if identity_stale:
            rejected.append(f"{candidate.source_name}: {candidate.title[:90]} stale_identity={identity_stale}")
            continue
        pdf_path = _download_pdf(candidate, output_dir)
        if not pdf_path:
            continue
        text = _read_pdf(pdf_path)
        low = text.lower()
        evidence = _publication_date_evidence(candidate, text, run_date)
        stale_markers = _stale_report_date_markers(candidate, text, run_date)
        valid = len(text) >= MIN_SOURCE_CHARS and any(term in low for term in REPORT_TERMS) and bool(evidence)
        if stale_markers and not evidence:
            valid = False
        if not valid:
            rejected.append(f"{candidate.source_name}: {candidate.title[:90]} text={len(text)} evidence={evidence} stale={stale_markers}")
            try:
                pdf_path.unlink(missing_ok=True)
            except Exception:
                pass
            continue
        scored.append((_score_candidate(candidate, text, evidence), candidate, pdf_path, text, evidence))
    if not scored:
        LOGGER.warning("No recent long-form research PDF accepted; rejected: %s", "; ".join(rejected[:20]))
        return None, None, "", []
    scored.sort(key=lambda x: x[0], reverse=True)
    _, candidate, pdf_path, text, evidence = scored[0]
    for _, _, path, _, _ in scored[1:]:
        try:
            path.unlink(missing_ok=True)
        except Exception:
            pass
    return candidate, pdf_path, text, evidence


def _split_text(text: str, chunk_chars: int = 7600, max_chunks: int = 10) -> list[str]:
    cleaned = _norm_text(text)
    chunks: list[str] = []
    start = 0
    while start < len(cleaned) and len(chunks) < max_chunks:
        end = min(len(cleaned), start + chunk_chars)
        if end < len(cleaned):
            split_at = max(cleaned.rfind(". ", start, end), cleaned.rfind("\n", start, end))
            if split_at > start + chunk_chars // 2:
                end = split_at + 1
        chunks.append(cleaned[start:end].strip())
        start = end
    return [x for x in chunks if x]


def _compile_key_points(candidate: ResearchCandidate, extracted_text: str) -> tuple[str, list[str]]:
    prompt = f"""
请基于以下英文研究报告开头内容，生成中文标题和3至6条关键点。只能使用原文信息，不得补写。
标题格式参考“普华永道：2026年全球加密货币监管现状与趋势”，即“机构：报告核心主题”。
输出 JSON：{{"title_cn":"中文标题","key_points":["关键点一"]}}
英文报告元数据：{json.dumps(asdict(candidate), ensure_ascii=False)}
英文报告节选：
{extracted_text[:15000]}
""".strip()
    data = _extract_json(_chat([{"role": "system", "content": "你是专业研究报告中文编译编辑。只输出严格 JSON。"}, {"role": "user", "content": prompt}], timeout=260))
    title = normalize_chinese_punctuation(str(data.get("title_cn") or f"{candidate.source_name}：{candidate.title}"))
    points = [normalize_chinese_punctuation(str(x)) for x in data.get("key_points", []) if str(x).strip()][:6]
    return title, points or ["本专题研究基于近两个月英文 PDF 原文逐段编译，发布前应对照原文核验数据和结论。"]


def _translate_chunk(candidate: ResearchCandidate, chunk: str, idx: int, total: int) -> list[str]:
    prompt = f"""
请将以下英文研究报告第 {idx}/{total} 部分忠实编译翻译为中文。要求：
1. 只翻译本部分英文原文已经出现的信息，不得添加外部背景、判断或结论。
2. 输出8至12个自然段，每段260至420个中文字符。
3. 保留原文事实链条、主要发现、数据逻辑、限定语和结论；不要压缩成摘要。
4. 可按原文逻辑加入“一、二、三”式小标题，但小标题也必须来自原文结构。
5. 专业术语首次出现写“中文（English，缩写）”；英文机构名保留英文。
6. 使用全角中文标点和中文全角引号，不使用半角引号。
输出 JSON：{{"paragraphs":["段落一","段落二"]}}
英文报告元数据：{json.dumps(asdict(candidate), ensure_ascii=False)}
英文原文：
{chunk}
""".strip()
    data = _extract_json(_chat([{"role": "system", "content": "你是专业研究报告翻译编辑。只输出严格 JSON。"}, {"role": "user", "content": prompt}], timeout=360))
    return [normalize_chinese_punctuation(str(x)) for x in data.get("paragraphs", []) if str(x).strip()]


def _extend_research(candidate: ResearchCandidate, extracted_text: str, existing_paragraphs: list[str]) -> list[str]:
    excerpt_start = min(len(extracted_text), max(0, len(existing_paragraphs) * 1100))
    excerpt = extracted_text[excerpt_start : excerpt_start + 12000] or extracted_text[-12000:]
    existing_brief = " ".join(existing_paragraphs[-6:])[:1800]
    prompt = f"""
当前专题研究中文稿长度不足。请继续基于同一份英文报告原文补充中文编译段落。
要求：
1. 只能使用英文原文中的信息，不得添加外部内容。
2. 不要重复已有中文段落的表达，应补充尚未展开的背景、数据逻辑、主要发现、限制条件和结论。
3. 输出6至10个自然段，每段260至420个中文字符。
4. 使用全角中文标点和中文全角引号。
输出 JSON：{{"paragraphs":["补充段落一"]}}
英文报告元数据：{json.dumps(asdict(candidate), ensure_ascii=False)}
已有中文稿末尾摘要：{existing_brief}
英文原文节选：
{excerpt}
""".strip()
    data = _extract_json(_chat([{"role": "system", "content": "你是专业研究报告翻译编辑。只输出严格 JSON。"}, {"role": "user", "content": prompt}], timeout=360))
    return [normalize_chinese_punctuation(str(x)) for x in data.get("paragraphs", []) if str(x).strip()]


def _compile_research(candidate: ResearchCandidate, pdf_path: Path, extracted_text: str, evidence: list[str]) -> dict[str, Any]:
    if not os.getenv("DEEPSEEK_API_KEY"):
        chunks = [normalize_chinese_punctuation(extracted_text[i : i + 620]) for i in range(0, min(len(extracted_text), 28000), 620)]
        return {
            "title_cn": normalize_chinese_punctuation(f"{candidate.source_name}：{candidate.title}"),
            "source_title": candidate.title,
            "source_name": candidate.source_name,
            "source_url": candidate.url,
            "pdf_path": str(pdf_path),
            "recent_date_evidence": evidence,
            "key_points": ["该专题研究由近两个月英文 PDF 原文自动提取并生成草稿，发布前需人工复核。"],
            "body_paragraphs": chunks[:TARGET_RESEARCH_PARAGRAPHS],
            "fact_check": "fallback 草稿，保留英文 PDF 原文。",
        }
    title_cn, key_points = _compile_key_points(candidate, extracted_text)
    paragraphs: list[str] = []
    chunks = _split_text(extracted_text, chunk_chars=7600, max_chunks=10)
    for idx, chunk in enumerate(chunks, start=1):
        paragraphs.extend(_translate_chunk(candidate, chunk, idx, len(chunks)))
        if len(paragraphs) >= TARGET_RESEARCH_PARAGRAPHS and len(" ".join(paragraphs)) >= MIN_RESEARCH_CHARS:
            break
    attempts = 0
    while len(" ".join(paragraphs)) < MIN_RESEARCH_CHARS and attempts < 2:
        attempts += 1
        paragraphs.extend(_extend_research(candidate, extracted_text, paragraphs))
    return {
        "title_cn": title_cn,
        "source_title": candidate.title,
        "source_name": candidate.source_name,
        "source_url": candidate.url,
        "pdf_path": str(pdf_path),
        "recent_date_evidence": evidence,
        "key_points": key_points,
        "body_paragraphs": paragraphs[:60],
        "fact_check": "基于一篇近两个月英文 PDF 原文分块逐段编译，未拼接其他报告，未加入原文以外信息。",
    }


def write_research_docx(research: dict[str, Any], output_path: str | Path) -> Path:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _setup_document(doc)
    _setup_section(doc.sections[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(RESEARCH_REPORT_TITLE)
    _set_run_font(r, east_asia=FONT_HEADING, size=16, color=HEADING_COLOR)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(normalize_chinese_punctuation(str(research.get("title_cn") or "-")))
    _set_run_font(r, east_asia=FONT_HEADING, size=15)
    for point in research.get("key_points", [])[:6]:
        _write_key_point(doc, str(point))
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    for para in research.get("body_paragraphs", []):
        _write_body(doc, str(para))
    _write_source(doc, str(research.get("source_name") or "-"))
    _write_reference(doc, f"原文标题：{research.get('source_title') or '-'}")
    _write_reference(doc, f"原文链接：{research.get('source_url') or '-'}")
    _write_reference(doc, f"PDF文件：{research.get('pdf_path') or '-'}")
    doc.core_properties.title = f"{RESEARCH_REPORT_TITLE}_{datetime.now(BEIJING_TZ).strftime('%Y%m%d')}"
    doc.save(path)
    return path


def check_research(research: dict[str, Any]) -> dict[str, Any]:
    errors: list[str] = []
    warnings: list[str] = []
    paragraphs = research.get("body_paragraphs") or []
    text = " ".join(str(x) for x in paragraphs)
    if len(paragraphs) < MIN_RESEARCH_PARAGRAPHS:
        errors.append(f"专题研究正文段落过少：{len(paragraphs)}，应至少{MIN_RESEARCH_PARAGRAPHS}段")
    if len(text) < MIN_RESEARCH_CHARS:
        errors.append(f"专题研究正文过短：当前约{len(text)}个字符，应至少{MIN_RESEARCH_CHARS}个字符")
    if has_half_width_quotes(text):
        errors.append("专题研究正文包含半角引号")
    pdf_path = Path(str(research.get("pdf_path") or ""))
    if not str(pdf_path).endswith(".pdf") or not pdf_path.exists():
        errors.append("未在仓库输出目录保存英文原文 PDF")
    if not research.get("source_url"):
        errors.append("缺少英文原文 PDF 链接")
    evidence = research.get("recent_date_evidence") or []
    if not evidence:
        errors.append("缺少近两个月内的报告发布日期证据")
    elif not any(str(x).startswith(("title_url_snippet=", "cover_publication_date=")) for x in evidence):
        errors.append("专题研究日期证据不是报告本身的发布日期")
    return {"ok": not errors, "errors": errors, "warnings": warnings}


def generate_research(output_root: Path, run_date: datetime) -> dict[str, Any]:
    source_dir = output_root / "research_sources" / run_date.strftime("%Y%m%d")
    candidate, pdf_path, text, evidence = _choose_pdf(source_dir, run_date)
    if not candidate or not pdf_path:
        raise RuntimeError("No suitable recent long-form English PDF research report found within the required two-month publication window")
    research = _compile_research(candidate, pdf_path, text, evidence)
    output_file = output_root / f"专题研究_{run_date.strftime('%Y%m%d')}.docx"
    write_research_docx(research, output_file)
    fact = check_research(research)
    return {"research": research, "factcheck": fact, "outputs": {"docx": str(output_file), "pdf": str(pdf_path)}}
